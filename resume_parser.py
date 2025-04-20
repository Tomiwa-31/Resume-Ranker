import os
import logging
from io import StringIO

# Configure logging
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

def parse_resume(file_path):
    """
    Parse a resume file and extract its text content.
    
    Args:
        file_path (str): Path to the resume file (PDF or DOCX)
        
    Returns:
        str: Extracted text content from the resume
    """
    logger.debug(f"Starting to parse resume: {file_path}")
    
    # Check if file exists
    if not os.path.exists(file_path):
        logger.error(f"File not found: {file_path}")
        raise FileNotFoundError(f"File not found: {file_path}")
        
    # Get file extension
    file_extension = file_path.split('.')[-1].lower()
    logger.debug(f"Detected file extension: {file_extension}")
    
    try:
        if file_extension == 'pdf':
            logger.debug("Processing as PDF")
            return parse_pdf(file_path)
        elif file_extension == 'docx':
            logger.debug("Processing as DOCX")
            return parse_docx(file_path)
        else:
            logger.error(f"Unsupported file format: {file_extension}")
            raise ValueError(f"Unsupported file format: {file_extension}")
    except Exception as e:
        logger.error(f"Error parsing resume: {str(e)}", exc_info=True)
        raise

def parse_pdf(file_path):
    """
    Parse PDF files using PyPDF2
    
    Args:
        file_path (str): Path to the PDF file
        
    Returns:
        str: Extracted text content
    """
    try:
        # Import PyPDF2 here to avoid unnecessary imports when parsing DOCX
        import PyPDF2
        logger.debug("PyPDF2 import successful")
        
        logger.debug(f"Opening PDF file: {file_path}")
        text = StringIO()
        with open(file_path, 'rb') as file:
            logger.debug("Creating PDF reader")
            pdf_reader = PyPDF2.PdfReader(file)
            
            # Log PDF info
            logger.debug(f"PDF pages: {len(pdf_reader.pages)}")
            
            # Extract text from each page
            for page_num in range(len(pdf_reader.pages)):
                logger.debug(f"Processing page {page_num+1}/{len(pdf_reader.pages)}")
                page = pdf_reader.pages[page_num]
                extracted_text = page.extract_text()
                text.write(extracted_text)
                
                # Log a sample of extracted text for debugging
                if page_num == 0:  # Only log first page
                    sample = extracted_text[:100] + "..." if len(extracted_text) > 100 else extracted_text
                    logger.debug(f"Sample text from page 1: {sample}")
        
        result = text.getvalue()
        logger.debug(f"Total text extracted: {len(result)} characters")
        return result
    except ImportError:
        logger.error("PyPDF2 is not installed. Please install it to parse PDF files.")
        raise
    except Exception as e:
        logger.error(f"Error parsing PDF: {str(e)}", exc_info=True)
        raise

def parse_docx(file_path):
    """
    Parse DOCX files using python-docx
    
    Args:
        file_path (str): Path to the DOCX file
        
    Returns:
        str: Extracted text content
    """
    try:
        # Import python-docx here to avoid unnecessary imports when parsing PDF
        import docx
        logger.debug("python-docx import successful")
        
        logger.debug(f"Opening DOCX file: {file_path}")
        doc = docx.Document(file_path)
        
        # Log document info
        logger.debug(f"Document paragraphs: {len(doc.paragraphs)}")
        logger.debug(f"Document tables: {len(doc.tables)}")
        
        # Extract text from paragraphs
        full_text = []
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                full_text.append(para.text)
                if i < 3:  # Log the first few paragraphs for debugging
                    logger.debug(f"Paragraph {i+1} text: {para.text[:50]}...")
            
        # Also extract text from tables
        table_text_count = 0
        for table_idx, table in enumerate(doc.tables):
            logger.debug(f"Processing table {table_idx+1}, rows: {len(table.rows)}")
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text.strip():
                            full_text.append(para.text)
                            table_text_count += 1
        
        logger.debug(f"Extracted {table_text_count} text elements from tables")
        
        result = '\n'.join(full_text)
        logger.debug(f"Total text extracted: {len(result)} characters")
        return result
    except ImportError:
        logger.error("python-docx is not installed. Please install it to parse DOCX files.")
        raise
    except Exception as e:
        logger.error(f"Error parsing DOCX: {str(e)}", exc_info=True)
        raise
